#!/usr/bin/env python3
"""
media_scan.py — Walk a directory (and subdirectories), find media files,
run ffprobe on each, and print a clean, colorized summary.

Usage:
    python media_scan.py "path/to/dir"
    python media_scan.py "path/to/dir" --save
    python media_scan.py "path/to/dir" -s -f ansi
    python media_scan.py "path/to/dir" -s -f md
    python media_scan.py "path/to/dir" -s -f docx
    python media_scan.py "path/to/dir" -s -o report.html

Requires: ffprobe (part of ffmpeg) available on PATH.
DOCX export additionally requires: pip install python-docx
"""

import os
import re
import sys
import json
import html
import shutil
import argparse
import subprocess
from datetime import timedelta, datetime


# --------------------------------------------------------------------------
# Terminal colors (ANSI) — always emitted; stripped at print-time if the
# destination isn't a real terminal, but kept intact for the save buffer.
# --------------------------------------------------------------------------
class C:
    RESET = "\033[0m"
    BOLD = "\033[1m"
    DIM = "\033[2m"
    ITALIC = "\033[3m"
    UNDER = "\033[4m"

    RED = "\033[31m"
    GREEN = "\033[32m"
    YELLOW = "\033[33m"
    BLUE = "\033[34m"
    MAGENTA = "\033[35m"
    CYAN = "\033[36m"
    WHITE = "\033[37m"

    B_RED = "\033[91m"
    B_GREEN = "\033[92m"
    B_YELLOW = "\033[93m"
    B_BLUE = "\033[94m"
    B_MAGENTA = "\033[95m"
    B_CYAN = "\033[96m"
    B_WHITE = "\033[97m"


def term_supports_color():
    if os.environ.get("NO_COLOR"):
        return False
    return sys.stdout.isatty() or bool(os.environ.get("FORCE_COLOR"))


TERM_COLOR = term_supports_color()
ANSI_RE = re.compile(r"\033\[[0-9;]*m")


def strip_ansi(text):
    return ANSI_RE.sub("", text)


def c(text, *codes):
    """Wrap text in ANSI codes. Always applied — stripped later at print
    time if stdout isn't a terminal, but kept for the save buffer."""
    if not codes:
        return str(text)
    return "".join(codes) + str(text) + C.RESET


# --------------------------------------------------------------------------
# Output buffering — every line goes through out(), which prints to the
# real terminal (color-stripped if needed) and, when --save is active,
# stores the full-color version for export.
# --------------------------------------------------------------------------
SAVE_BUFFER = None  # set to a list when --save is requested


def out(line=""):
    if SAVE_BUFFER is not None:
        SAVE_BUFFER.append(line)
    print(line if TERM_COLOR else strip_ansi(line))


# --------------------------------------------------------------------------
# Media file extensions
# --------------------------------------------------------------------------
VIDEO_EXTS = {
    ".mp4",
    ".mkv",
    ".mov",
    ".avi",
    ".webm",
    ".flv",
    ".wmv",
    ".m4v",
    ".mpg",
    ".mpeg",
    ".3gp",
    ".ts",
    ".m2ts",
    ".vob",
    ".ogv",
}
AUDIO_EXTS = {
    ".mp3",
    ".wav",
    ".flac",
    ".aac",
    ".ogg",
    ".m4a",
    ".wma",
    ".opus",
    ".aiff",
    ".alac",
}
MEDIA_EXTS = VIDEO_EXTS | AUDIO_EXTS


# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------
def human_size(num_bytes):
    """Binary (1024-based) human readable size, like KiB/MiB/GiB."""
    if num_bytes is None:
        return "N/A"
    size = float(num_bytes)
    for unit in ["B", "KiB", "MiB", "GiB", "TiB", "PiB"]:
        if size < 1024.0:
            return f"{size:.2f} {unit}"
        size /= 1024.0
    return f"{size:.2f} EiB"


def human_bitrate(bps):
    """Human readable bitrate (decimal-based, standard broadcast convention)."""
    if bps is None:
        return "N/A"
    bps = float(bps)
    for unit in ["bps", "kbps", "Mbps", "Gbps"]:
        if bps < 1000.0:
            return f"{bps:.0f} {unit}"
        bps /= 1000.0
    return f"{bps:.2f} Tbps"


def format_timecode(total_frames, fps):
    """H:MM:SS:FF timecode from an exact frame count + frame rate."""
    fps_int = max(1, round(fps))
    total_frames = int(round(total_frames))
    hours, rem = divmod(total_frames, fps_int * 3600)
    minutes, rem = divmod(rem, fps_int * 60)
    secs, frames = divmod(rem, fps_int)
    return f"{hours}:{minutes:02d}:{secs:02d}:{frames:02d}"


def human_duration(seconds, fps=None, frame_count=None):
    """Frame-accurate duration for video (H:MM:SS:FF), millisecond
    precision for audio/no-fps sources — avoids the old whole-second
    rounding (e.g. 7s15f no longer gets rounded up to 8s)."""
    if seconds is None and frame_count is None:
        return "N/A"

    if fps:
        if frame_count:
            # Exact frame count reported by ffprobe — most accurate.
            return f"{format_timecode(frame_count, fps)}  ({frame_count}f)"
        if seconds is not None:
            # No exact frame count available — estimate from duration*fps.
            est_frames = round(float(seconds) * fps)
            return f"{format_timecode(est_frames, fps)}  (~{est_frames}f)"

    try:
        seconds = float(seconds)
    except (TypeError, ValueError):
        return "N/A"
    total_seconds = int(seconds)
    hours, rem = divmod(total_seconds, 3600)
    minutes, secs = divmod(rem, 60)
    millis = round((seconds - total_seconds) * 1000)
    if millis == 1000:
        millis = 0
        secs += 1
    return f"{hours}:{minutes:02d}:{secs:02d}.{millis:03d}"


def parse_fps(rate_str):
    """ffprobe returns fps as a fraction string like '30000/1001'."""
    if not rate_str or rate_str == "0/0":
        return None
    try:
        if "/" in rate_str:
            num, den = rate_str.split("/")
            num, den = float(num), float(den)
            if den == 0:
                return None
            return round(num / den, 3)
        return round(float(rate_str), 3)
    except (ValueError, ZeroDivisionError):
        return None


def run_ffprobe(filepath):
    """Run ffprobe and return parsed JSON, or None on failure."""
    cmd = [
        "ffprobe",
        "-v",
        "quiet",
        "-print_format",
        "json",
        "-show_format",
        "-show_streams",
        filepath,
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode != 0:
            return None, result.stderr.strip()
        return json.loads(result.stdout), None
    except FileNotFoundError:
        out(c("Error: ffprobe not found on PATH. Install ffmpeg first.", C.B_RED, C.BOLD))
        sys.exit(1)
    except subprocess.TimeoutExpired:
        return None, "ffprobe timed out"
    except json.JSONDecodeError:
        return None, "failed to parse ffprobe output"


def extract_info(filepath):
    """Return a dict of extracted media info, or an error dict."""
    data, err = run_ffprobe(filepath)
    if data is None:
        return {"error": err or "unknown ffprobe error"}

    fmt = data.get("format", {}) or {}
    streams = data.get("streams", []) or []

    video_stream = next((s for s in streams if s.get("codec_type") == "video"), None)
    audio_stream = next((s for s in streams if s.get("codec_type") == "audio"), None)

    info = {}

    try:
        info["size_bytes"] = os.path.getsize(filepath)
    except OSError:
        info["size_bytes"] = fmt.get("size")

    duration = fmt.get("duration")
    if duration is None and video_stream:
        duration = video_stream.get("duration")
    if duration is None and audio_stream:
        duration = audio_stream.get("duration")
    info["duration"] = duration

    info["bitrate"] = fmt.get("bit_rate")

    if video_stream:
        info["kind"] = "video"
        info["codec"] = video_stream.get("codec_name", "unknown").upper()
        info["codec_long"] = video_stream.get("codec_long_name", "")
        info["width"] = video_stream.get("width")
        info["height"] = video_stream.get("height")

        fps = parse_fps(video_stream.get("avg_frame_rate")) or parse_fps(video_stream.get("r_frame_rate"))
        info["fps"] = fps

        info["video_bitrate"] = video_stream.get("bit_rate")
        info["pix_fmt"] = video_stream.get("pix_fmt")

        # Exact frame count, when the container/ffprobe reports it —
        # gives frame-accurate duration instead of an estimate.
        nb_frames_raw = video_stream.get("nb_frames")
        try:
            info["frame_count"] = int(nb_frames_raw) if nb_frames_raw else None
        except (TypeError, ValueError):
            info["frame_count"] = None

        if audio_stream:
            info["audio_codec"] = audio_stream.get("codec_name", "unknown").upper()
            info["audio_channels"] = audio_stream.get("channels")
            info["audio_sample_rate"] = audio_stream.get("sample_rate")
    elif audio_stream:
        info["kind"] = "audio"
        info["codec"] = audio_stream.get("codec_name", "unknown").upper()
        info["codec_long"] = audio_stream.get("codec_long_name", "")
        info["audio_channels"] = audio_stream.get("channels")
        info["audio_sample_rate"] = audio_stream.get("sample_rate")
        info["audio_bitrate"] = audio_stream.get("bit_rate")
    else:
        return {"error": "no audio or video stream found"}

    return info


# --------------------------------------------------------------------------
# Output formatting
# --------------------------------------------------------------------------
def term_width():
    return min(shutil.get_terminal_size((80, 20)).columns, 100)


def print_header(title):
    bar = "─" * term_width()
    out()
    out(c(bar, C.B_BLUE))
    out(c(f"  {title}", C.BOLD, C.B_CYAN))
    out(c(bar, C.B_BLUE))


def print_dir_header(path, count):
    out()
    out(c(f"📁 {path}", C.BOLD, C.B_YELLOW) + c(f"  ({count} media file{'s' if count != 1 else ''})", C.DIM))
    out(c("─" * term_width(), C.DIM))


def info_rows(info):
    """Return [(label, value_str), ...] for a media info dict — the single
    source of truth used by terminal, Markdown, and DOCX output so they
    never drift out of sync."""
    rows = []
    if info.get("kind") == "video":
        codec = info.get("codec", "N/A") + (f" ({info['codec_long']})" if info.get("codec_long") else "")
        rows.append(("Codec", codec))
        if info.get("width") and info.get("height"):
            rows.append(("Dimensions", f"{info['width']}x{info['height']}"))
        rows.append(("FPS", f"{info['fps']:.2f}" if info.get("fps") else "N/A"))
        bitrate = info.get("video_bitrate") or info.get("bitrate")
        rows.append(("Bitrate", human_bitrate(bitrate)))
        rows.append(
            ("Duration", human_duration(info.get("duration"), fps=info.get("fps"), frame_count=info.get("frame_count")))
        )
        rows.append(("Size", human_size(info.get("size_bytes"))))
        if info.get("audio_codec"):
            rows.append(
                (
                    "Audio",
                    f"{info['audio_codec']} · {info.get('audio_channels', '?')}ch · "
                    f"{info.get('audio_sample_rate', '?')} Hz",
                )
            )
    else:  # audio
        codec = info.get("codec", "N/A") + (f" ({info['codec_long']})" if info.get("codec_long") else "")
        rows.append(("Codec", codec))
        rows.append(("Bitrate", human_bitrate(info.get("audio_bitrate") or info.get("bitrate"))))
        rows.append(("Sample Rate", f"{info.get('audio_sample_rate', 'N/A')} Hz"))
        rows.append(("Channels", str(info.get("audio_channels", "N/A"))))
        rows.append(("Duration", human_duration(info.get("duration"))))
        rows.append(("Size", human_size(info.get("size_bytes"))))
    return rows


def print_file_info(filename, info):
    kind_icon = "🎬" if info.get("kind") == "video" else "🎵"
    out()
    out(c(f"  {kind_icon} [{filename}]", C.BOLD, C.B_MAGENTA))

    if "error" in info:
        out(c(f"      ⚠ Could not read media info: {info['error']}", C.RED))
        return

    for label, value in info_rows(info):
        out(f"      {c(label + ':', C.B_GREEN):<28} {c(value, C.WHITE)}")


def print_summary(total_files, total_size, error_count):
    bar = "─" * term_width()
    out()
    out(c(bar, C.B_BLUE))
    out(c("  Summary", C.BOLD, C.B_CYAN))
    out(f"    {c('Files scanned:', C.B_GREEN)} {total_files}")
    out(f"    {c('Total size:', C.B_GREEN)} {human_size(total_size)}")
    if error_count:
        out(f"    {c('Errors:', C.B_RED)} {error_count}")
    out(c(bar, C.B_BLUE))
    out()


# --------------------------------------------------------------------------
# ANSI -> HTML conversion (for --save --format html)
# --------------------------------------------------------------------------
_COLOR_MAP = {
    30: "#6b6b6b",
    31: "#ff5555",
    32: "#50fa7b",
    33: "#f1fa8c",
    34: "#6c93ff",
    35: "#ff79c6",
    36: "#8be9fd",
    37: "#f8f8f2",
    90: "#7b7b7b",
    91: "#ff6e6e",
    92: "#69ff94",
    93: "#ffffa5",
    94: "#8aa6ff",
    95: "#ff92df",
    96: "#a4ffff",
    97: "#ffffff",
}


def ansi_to_html(text):
    """Convert a string containing ANSI escape codes into HTML spans."""
    parts = []
    state = {"bold": False, "dim": False, "underline": False, "color": None}
    pos = 0

    def style_str():
        s = []
        if state["bold"]:
            s.append("font-weight:bold")
        if state["dim"]:
            s.append("opacity:0.6")
        if state["underline"]:
            s.append("text-decoration:underline")
        if state["color"]:
            s.append(f"color:{state['color']}")
        return ";".join(s)

    for m in ANSI_RE.finditer(text):
        chunk = text[pos : m.start()]
        if chunk:
            st = style_str()
            if st:
                parts.append(f'<span style="{st}">{html.escape(chunk)}</span>')
            else:
                parts.append(html.escape(chunk))
        codes_raw = m.group(0)[2:-1]  # strip \033[ and m
        codes_list = [int(x) for x in codes_raw.split(";") if x != ""] or [0]
        for code in codes_list:
            if code == 0:
                state = {"bold": False, "dim": False, "underline": False, "color": None}
            elif code == 1:
                state["bold"] = True
            elif code == 2:
                state["dim"] = True
            elif code == 4:
                state["underline"] = True
            elif code in _COLOR_MAP:
                state["color"] = _COLOR_MAP[code]
        pos = m.end()

    chunk = text[pos:]
    if chunk:
        st = style_str()
        if st:
            parts.append(f'<span style="{st}">{html.escape(chunk)}</span>')
        else:
            parts.append(html.escape(chunk))

    return "".join(parts)


HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{
    background: #1e1f29;
    margin: 0;
    padding: 2rem;
    display: flex;
    justify-content: center;
  }}
  pre.terminal {{
    background: #282a36;
    color: #f8f8f2;
    font-family: 'SF Mono', 'Fira Code', Consolas, 'Courier New', monospace;
    font-size: 14px;
    line-height: 1.5;
    padding: 1.5rem 2rem;
    border-radius: 10px;
    box-shadow: 0 8px 30px rgba(0,0,0,0.4);
    max-width: 1000px;
    width: 100%;
    overflow-x: auto;
    white-space: pre-wrap;
    word-break: break-word;
  }}
  .meta {{
    color: #6272a4;
    font-family: -apple-system, sans-serif;
    font-size: 12px;
    max-width: 1000px;
    width: 100%;
    margin: 0 auto 0.5rem auto;
    text-align: right;
  }}
</style>
</head>
<body>
<div style="width:100%;max-width:1000px;margin:0 auto;">
  <div class="meta">Generated {timestamp}</div>
  <pre class="terminal">{body}</pre>
</div>
</body>
</html>
"""


def save_report(fmt, output_path, root_path, buffer_lines, scan_results, total_files, total_size, error_count):
    if fmt in ("html", "ansi"):
        full_text = "\n".join(buffer_lines)
        if fmt == "ansi":
            content = full_text
        else:
            body_html = ansi_to_html(full_text)
            content = HTML_TEMPLATE.format(
                title=f"Media Scan Report — {os.path.basename(os.path.abspath(root_path))}",
                timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                body=body_html,
            )
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)

    elif fmt == "md":
        content = build_markdown_report(scan_results, total_files, total_size, error_count, root_path)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)

    elif fmt == "docx":
        build_docx_report(scan_results, total_files, total_size, error_count, root_path, output_path)


# --------------------------------------------------------------------------
# Markdown export — no native color support in Markdown, so the "clean"
# look comes from structure instead: headings + tables render nicely on
# GitHub, GitLab, Notion, most forums/wikis that support standard MD.
# --------------------------------------------------------------------------
def build_markdown_report(scan_results, total_files, total_size, error_count, root_path):
    lines = [
        "# 🎞️ Media Scan Report",
        "",
        f"**Directory:** `{os.path.abspath(root_path)}`  ",
        f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "## Summary",
        "",
        "| Metric | Value |",
        "|---|---|",
        f"| Files scanned | {total_files} |",
        f"| Total size | {human_size(total_size)} |",
    ]
    if error_count:
        lines.append(f"| Errors | {error_count} |")
    lines.append("")

    for entry in scan_results:
        count = len(entry["files"])
        lines.append(f"## 📁 {entry['dir']} ({count} file{'s' if count != 1 else ''})")
        lines.append("")
        for filename, info in entry["files"]:
            icon = "🎬" if info.get("kind") == "video" else ("🎵" if info.get("kind") == "audio" else "⚠️")
            lines.append(f"### {icon} `{filename}`")
            lines.append("")
            if "error" in info:
                lines.append(f"> ⚠ Could not read media info: {info['error']}")
                lines.append("")
                continue
            lines.append("| Field | Value |")
            lines.append("|---|---|")
            for label, value in info_rows(info):
                lines.append(f"| {label} | {str(value).replace('|', chr(92) + '|')} |")
            lines.append("")

    return "\n".join(lines)


# --------------------------------------------------------------------------
# DOCX export — uses python-docx. Real colored text via font.color.rgb,
# so unlike Markdown this keeps genuine color coding (green labels, red
# errors, etc.) inside a normal Word document.
# --------------------------------------------------------------------------
DOCX_GREEN = (0x1E, 0x7A, 0x34)
DOCX_RED = (0xC0, 0x1C, 0x28)
DOCX_GRAY = (0x66, 0x66, 0x66)
DOCX_BLUE = (0x1F, 0x4E, 0x9C)


def build_docx_report(scan_results, total_files, total_size, error_count, root_path, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        print("Error: the 'python-docx' package is required for --format docx.")
        print("Install it with: pip install python-docx")
        sys.exit(1)

    doc = Document()

    doc.add_heading("🎞 Media Scan Report", level=0)

    p = doc.add_paragraph()
    p.add_run("Directory: ").bold = True
    p.add_run(os.path.abspath(root_path))
    p = doc.add_paragraph()
    p.add_run("Generated: ").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    doc.add_heading("Summary", level=1)
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Light Grid Accent 1"
    hdr = summary_table.rows[0].cells
    hdr[0].text, hdr[1].text = "Metric", "Value"
    summary_rows = [("Files scanned", str(total_files)), ("Total size", human_size(total_size))]
    if error_count:
        summary_rows.append(("Errors", str(error_count)))
    for label, value in summary_rows:
        row = summary_table.add_row().cells
        row[0].text, row[1].text = label, value

    for entry in scan_results:
        count = len(entry["files"])
        doc.add_heading(f"📁 {entry['dir']}  ({count} file{'s' if count != 1 else ''})", level=1)

        for filename, info in entry["files"]:
            icon = "🎬" if info.get("kind") == "video" else "🎵"
            doc.add_heading(f"{icon} {filename}", level=2)

            if "error" in info:
                run = doc.add_paragraph().add_run(f"⚠ Could not read media info: {info['error']}")
                run.font.color.rgb = RGBColor(*DOCX_RED)
                continue

            table = doc.add_table(rows=0, cols=2)
            table.style = "Light List Accent 1"
            for label, value in info_rows(info):
                cells = table.add_row().cells
                label_run = cells[0].paragraphs[0].add_run(label)
                label_run.bold = True
                label_run.font.color.rgb = RGBColor(*DOCX_GREEN)
                cells[1].text = str(value)

            doc.add_paragraph()  # spacing between files

    doc.save(output_path)


# --------------------------------------------------------------------------
# Main
# --------------------------------------------------------------------------
def is_media_file(filename):
    ext = os.path.splitext(filename)[1].lower()
    return ext in MEDIA_EXTS


def build_arg_parser():
    p = argparse.ArgumentParser(description="Scan a directory for media files and report their ffprobe stats.")
    p.add_argument("path", help="Directory to scan")
    p.add_argument("-s", "--save", action="store_true", help="Save the report to a file in addition to printing it")
    p.add_argument(
        "-f",
        "--format",
        choices=["html", "ansi", "md", "docx"],
        default="html",
        help="Save format: 'html' (colored, viewable in any browser — good for "
        "Teams/Slack/forum sharing), 'ansi' (raw text with terminal color "
        "codes), 'md' (Markdown with tables — for GitHub/GitLab/wikis), or "
        "'docx' (Word document with real colored text; requires the "
        "python-docx package). Default: html",
    )
    p.add_argument(
        "-o",
        "--output",
        default=None,
        help="Output file path (default: media_scan_report_<timestamp>.<ext> in the cwd)",
    )
    return p


def main():
    global SAVE_BUFFER

    args = build_arg_parser().parse_args()
    root_path = args.path

    if not os.path.isdir(root_path):
        print(
            c(f"Error: '{root_path}' is not a valid directory.", C.B_RED, C.BOLD)
            if TERM_COLOR
            else f"Error: '{root_path}' is not a valid directory."
        )
        sys.exit(1)

    if shutil.which("ffprobe") is None:
        print("Error: ffprobe not found on PATH. Install ffmpeg first.")
        sys.exit(1)

    if args.save:
        SAVE_BUFFER = []

    print_header(f"Media Scan: {os.path.abspath(root_path)}")

    total_files = 0
    total_size = 0
    error_count = 0
    scan_results = []  # structured data for md/docx export: [{"dir": ..., "files": [(name, info), ...]}]

    for dirpath, dirnames, filenames in os.walk(root_path):
        dirnames.sort()
        media_files = sorted(f for f in filenames if is_media_file(f))
        if not media_files:
            continue

        rel_dir = os.path.relpath(dirpath, root_path)
        display_dir = "." if rel_dir == "." else rel_dir
        print_dir_header(display_dir, len(media_files))

        dir_entry = {"dir": display_dir, "files": []}

        for filename in media_files:
            filepath = os.path.join(dirpath, filename)
            info = extract_info(filepath)
            print_file_info(filename, info)
            dir_entry["files"].append((filename, info))

            total_files += 1
            if "error" in info:
                error_count += 1
            else:
                total_size += info.get("size_bytes") or 0  # type: ignore

        scan_results.append(dir_entry)

    if total_files == 0:
        out()
        out(c("  No media files found.", C.B_YELLOW))
        out()
    else:
        print_summary(total_files, total_size, error_count)

    if args.save:
        ext_map = {"html": "html", "ansi": "ans", "md": "md", "docx": "docx"}
        ext = ext_map[args.format]
        if args.output:
            output_path = args.output
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"temp\\media_scan_report_{timestamp}.{ext}"

        save_report(
            args.format, output_path, root_path, SAVE_BUFFER, scan_results, total_files, total_size, error_count
        )
        print(
            c(f"✔ Report saved to {output_path}", C.B_GREEN, C.BOLD) if TERM_COLOR else f"Report saved to {output_path}"
        )


if __name__ == "__main__":
    main()
