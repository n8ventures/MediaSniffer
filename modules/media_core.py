"""
media_core.py — Non-GUI logic for the media scanner: binary resolution,
ffprobe/ffmpeg wrappers, metadata + loudness extraction, and report builders
(HTML / Markdown / DOCX / TXT). Kept separate from the GUI so it's easy to
test, reuse from a CLI, or swap pieces out.
"""

import os
import re
import sys
import json
import html
import shutil
import subprocess
from datetime import datetime

# --------------------------------------------------------------------------
# File types — GIF + video only (no plain audio files)
# --------------------------------------------------------------------------
VIDEO_EXTS = {
    ".mp4",
    ".mkv",
    ".mov",
    ".avi",
    ".webm",
    ".flv",
    ".wmv",
    ".m4v",
    ".mpg",
    ".mpeg",
    ".3gp",
    ".ts",
    ".m2ts",
    ".vob",
    ".ogv",
}
GIF_EXTS = {".gif"}
MEDIA_EXTS = VIDEO_EXTS | GIF_EXTS


def is_media_file(filename):
    return os.path.splitext(filename)[1].lower() in MEDIA_EXTS


# --------------------------------------------------------------------------
# Binary resolution — ffmpeg/ffprobe are fetched once at *build* time (see
# devtools/fetch_ffmpeg.py or however you're pulling them) into:
#   bin/Silicon/   macOS arm64  — from ffmpeg.martin-riedl.de
#   bin/Win64/     Windows x64  — from gyan.dev
# and bundled straight into the app by the .spec, so there's no runtime
# download and no static_ffmpeg dependency. Only macOS Apple Silicon and
# Windows x64 are covered by a bundled copy right now — Intel Mac and Linux
# fall back to whatever's on PATH.
# --------------------------------------------------------------------------
from modules.platformModules import win, mac, bundle_path

_FFMPEG_PATH = None
_FFPROBE_PATH = None


def set_binaries(ffmpeg_path, ffprobe_path):
    """Explicitly set the resolved ffmpeg/ffprobe paths, bypassing the
    lookup below entirely on every future resolve_binaries() call."""
    global _FFMPEG_PATH, _FFPROBE_PATH
    _FFMPEG_PATH, _FFPROBE_PATH = ffmpeg_path, ffprobe_path


def _bundled_bin_candidates():
    """Candidate roots to search for bin/<Silicon|Win64>/, covering both
    running from source (project root) and the built app. PyInstaller's
    macOS BUNDLE step doesn't always place non-.dylib binaries in the same
    spot as `datas` (Resources vs. Frameworks depending on version/how
    they're declared in the .spec) — check a couple of likely roots rather
    than assuming one. Verify with `find *.app -name ffmpeg` after a build
    and trim this list down once you know where yours actually lands.
    """
    folder = "Win64" if win else "Silicon" if mac else None
    if not folder:
        return []

    roots = []
    if bundle_path:
        roots.append(bundle_path)  # Contents/Resources on macOS
        roots.append(os.path.join(os.path.dirname(bundle_path), "Frameworks"))
        roots.append(os.path.dirname(sys.executable))
    else:
        roots.append(".")  # running from source, project root
    return [os.path.join(root, "bin", folder) for root in roots]


def resolve_binaries():
    """Resolve (ffmpeg_path, ffprobe_path), caching the result."""
    global _FFMPEG_PATH, _FFPROBE_PATH
    if _FFMPEG_PATH and _FFPROBE_PATH:
        return _FFMPEG_PATH, _FFPROBE_PATH

    ext = ".exe" if win else ""
    for bin_dir in _bundled_bin_candidates():
        ffmpeg_path = os.path.join(bin_dir, f"ffmpeg{ext}")
        ffprobe_path = os.path.join(bin_dir, f"ffprobe{ext}")
        if os.path.isfile(ffmpeg_path) and os.path.isfile(ffprobe_path):
            if not win:
                for exe in (ffmpeg_path, ffprobe_path):
                    if not os.access(exe, os.X_OK):
                        os.chmod(exe, os.stat(exe).st_mode | 0o111)
            _FFMPEG_PATH, _FFPROBE_PATH = ffmpeg_path, ffprobe_path
            return _FFMPEG_PATH, _FFPROBE_PATH

    ffmpeg_path = shutil.which("ffmpeg")
    ffprobe_path = shutil.which("ffprobe")
    _FFMPEG_PATH, _FFPROBE_PATH = ffmpeg_path, ffprobe_path
    return ffmpeg_path, ffprobe_path


def binaries_available():
    ffmpeg_path, ffprobe_path = resolve_binaries()
    return bool(ffmpeg_path and ffprobe_path)


# --------------------------------------------------------------------------
# Formatting helpers
# --------------------------------------------------------------------------
def human_size(num_bytes):
    """Binary (1024-based): KiB/MiB/GiB, not the decimal 1000-based kind."""
    if num_bytes is None:
        return "N/A"
    size = float(num_bytes)
    for unit in ["B", "KiB", "MiB", "GiB", "TiB", "PiB"]:
        if size < 1024.0:
            return f"{size:.2f} {unit}"
        size /= 1024.0
    return f"{size:.2f} EiB"


def human_bitrate(bps):
    if bps is None:
        return "N/A"
    bps = float(bps)
    for unit in ["bps", "kbps", "Mbps", "Gbps"]:
        if bps < 1000.0:
            return f"{bps:.0f} {unit}"
        bps /= 1000.0
    return f"{bps:.2f} Tbps"


def parse_fps(rate_str):
    if not rate_str or rate_str == "0/0":
        return None
    try:
        if "/" in rate_str:
            num, den = rate_str.split("/")
            num, den = float(num), float(den)
            return round(num / den, 3) if den else None
        return round(float(rate_str), 3)
    except (ValueError, ZeroDivisionError):
        return None


def format_timecode(total_frames, fps):
    """H:MM:SS:FF timecode from an exact frame count + frame rate."""
    fps_int = max(1, round(fps))
    total_frames = int(round(total_frames))
    hours, rem = divmod(total_frames, fps_int * 3600)
    minutes, rem = divmod(rem, fps_int * 60)
    secs, frames = divmod(rem, fps_int)
    return f"{hours}:{minutes:02d}:{secs:02d}:{frames:02d}"


def human_duration(seconds, fps=None, frame_count=None):
    """Frame-accurate H:MM:SS:FF for video/gif — no whole-second rounding."""
    if seconds is None and frame_count is None:
        return "N/A"
    if fps:
        if frame_count:
            return f"{format_timecode(frame_count, fps)}  ({frame_count}f)"
        if seconds is not None:
            est = round(float(seconds) * fps)
            return f"{format_timecode(est, fps)}  (~{est}f)"
    try:
        seconds = float(seconds)
    except (TypeError, ValueError):
        return "N/A"
    total_seconds = int(seconds)
    hours, rem = divmod(total_seconds, 3600)
    minutes, secs = divmod(rem, 60)
    millis = round((seconds - total_seconds) * 1000)
    if millis == 1000:
        millis = 0
        secs += 1
    return f"{hours}:{minutes:02d}:{secs:02d}.{millis:03d}"


def bit_depth_from_pix_fmt(pix_fmt):
    if not pix_fmt:
        return None
    m = re.search(r"(\d+)(?:le|be)?$", pix_fmt)
    if m and "p" in pix_fmt:
        return int(m.group(1))
    return 8  # yuv420p etc. with no suffix is 8-bit


def aspect_ratio_string(width, height, dar=None):
    if dar and dar not in ("0:1", "N/A"):
        return dar.replace(":", ":")
    if not width or not height:
        return "N/A"
    import math

    g = math.gcd(width, height)
    return f"{width // g}:{height // g}"


# ProRes' `profile` field reports plain names (Proxy/LT/Standard/HQ/4444/XQ);
# map to Apple's actual marketing names for display.
_PRORES_PROFILE_MAP = {
    "proxy": "422 Proxy",
    "lt": "422 LT",
    "standard": "422",
    "hq": "422 HQ",
    "4444": "4444",
    "xq": "4444 XQ",
}


def format_codec_profile(codec_name, profile, level):
    """Human-readable profile/level suffix for the Codec field. ffprobe
    exposes this the same way for ProRes, H.264, HEVC, and DNxHD/DNxHR —
    codec_name alone is generic (e.g. always "prores" regardless of which
    of the 6 ProRes variants it is); the actual variant lives in `profile`
    (and `level`, meaningful only for H.264/HEVC)."""
    if not profile or str(profile).strip().lower() in ("unknown", "", "none"):
        return None

    codec_name = (codec_name or "").lower()
    profile = str(profile)

    if codec_name == "prores":
        return "ProRes " + _PRORES_PROFILE_MAP.get(profile.lower(), profile)

    if codec_name in ("h264", "hevc", "h265"):
        label = f"{profile} Profile"
        try:
            level = int(level)
            if level > 0:
                label += f" @ L{level / 10:.1f}"
        except (TypeError, ValueError):
            pass
        return label

    if codec_name == "dnxhd":
        # ffprobe reports e.g. "DNXHD" or "DNXHR HQX" — normalize casing
        return profile.replace("DNXHR", "DNxHR").replace("DNXHD", "DNxHD")

    return profile  # generic fallback: whatever ffprobe called it


# --------------------------------------------------------------------------
# ffprobe — required + optional metadata (single call per file)
# --------------------------------------------------------------------------
def run_ffprobe(filepath):
    _, ffprobe_path = resolve_binaries()
    if not ffprobe_path:
        return None, "ffprobe not available"
    cmd = [
        ffprobe_path,
        "-v",
        "quiet",
        "-print_format",
        "json",
        "-show_format",
        "-show_streams",
        filepath,
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if result.returncode != 0:
            return None, result.stderr.strip()
        return json.loads(result.stdout), None
    except subprocess.TimeoutExpired:
        return None, "ffprobe timed out"
    except json.JSONDecodeError:
        return None, "failed to parse ffprobe output"


def extract_info(filepath, options):
    """Return a metadata dict for one file. `options` is a dict of booleans
    keyed by the GUI checkbox ids (see gui_app.CHECKBOX_DEFS)."""
    data, err = run_ffprobe(filepath)
    if data is None:
        return {"error": err or "unknown ffprobe error"}

    fmt = data.get("format", {}) or {}
    streams = data.get("streams", []) or []
    video_stream = next((s for s in streams if s.get("codec_type") == "video"), None)
    audio_stream = next((s for s in streams if s.get("codec_type") == "audio"), None)

    if video_stream is None:
        return {"error": "no video/gif stream found"}

    info = {"kind": "gif" if filepath.lower().endswith(".gif") else "video"}

    try:
        info["size_bytes"] = os.path.getsize(filepath)
    except OSError:
        info["size_bytes"] = fmt.get("size")

    # --- Required fields ---
    info["codec"] = video_stream.get("codec_name", "unknown").upper()
    info["codec_long"] = video_stream.get("codec_long_name", "")
    info["codec_profile"] = format_codec_profile(
        video_stream.get("codec_name"), video_stream.get("profile"), video_stream.get("level")
    )
    info["codec_tag"] = video_stream.get("codec_tag_string", "").strip() or None
    info["width"] = video_stream.get("width")
    info["height"] = video_stream.get("height")
    fps = parse_fps(video_stream.get("avg_frame_rate")) or parse_fps(video_stream.get("r_frame_rate"))
    info["fps"] = fps

    duration = fmt.get("duration") or video_stream.get("duration")
    info["duration"] = duration
    nb_frames_raw = video_stream.get("nb_frames")
    try:
        info["frame_count"] = int(nb_frames_raw) if nb_frames_raw else None
    except (TypeError, ValueError):
        info["frame_count"] = None

    info["bitrate"] = video_stream.get("bit_rate") or fmt.get("bit_rate")

    # --- Audio presence (needed regardless of checkboxes, to know whether
    #     loudness analysis is even possible) ---
    info["has_audio"] = audio_stream is not None

    # --- Optional: audio track info ---
    if options.get("audio_info") and audio_stream:
        info["audio_codec"] = audio_stream.get("codec_name", "unknown").upper()
        info["audio_channels"] = audio_stream.get("channels")
        info["audio_sample_rate"] = audio_stream.get("sample_rate")
        info["audio_bitrate"] = audio_stream.get("bit_rate")

    # --- Optional: color / HDR info ---
    if options.get("color_info"):
        pix_fmt = video_stream.get("pix_fmt")
        transfer = video_stream.get("color_transfer", "")
        is_hdr = transfer in ("smpte2084", "arib-std-b67")
        hdr_label = "HDR10 (PQ)" if transfer == "smpte2084" else "HLG" if transfer == "arib-std-b67" else "SDR"
        depth = bit_depth_from_pix_fmt(pix_fmt)
        primaries = video_stream.get("color_primaries") or "N/A"
        info["color_summary"] = f"{depth}-bit · {primaries} · {hdr_label}" if depth else "N/A"
        info["is_hdr"] = is_hdr

    # --- Optional: container/format name ---
    if options.get("container_info"):
        info["container"] = fmt.get("format_long_name") or fmt.get("format_name") or "N/A"

    # --- Optional: creation date ---
    if options.get("creation_date"):
        tags = fmt.get("tags", {}) or {}
        vtags = video_stream.get("tags", {}) or {}
        info["creation_date"] = tags.get("creation_time") or vtags.get("creation_time") or "N/A"

    # --- Optional: aspect ratio ---
    if options.get("aspect_ratio"):
        info["aspect_ratio"] = aspect_ratio_string(
            info["width"], info["height"], video_stream.get("display_aspect_ratio")
        )

    # --- Optional: loudness (LUFS / true peak) — only meaningful w/ audio ---
    if options.get("lufs") or options.get("true_peak"):
        if not info["has_audio"]:
            if options.get("lufs"):
                info["lufs_integrated"] = None
            if options.get("true_peak"):
                info["true_peak_db"] = None
        else:
            loud = analyze_loudness(filepath, options.get("lufs", False), options.get("true_peak", False))
            info.update(loud)  # type: ignore

    # --- Optional: TVC slate beep (only fires when the heuristic matches) ---
    if options.get("tvc_slate"):
        info["is_tvc_candidate"] = is_probable_tvc(info, filepath)
        if info["is_tvc_candidate"] and info["has_audio"]:
            slate = analyze_slate_beep(filepath)
            info["slate_beep_peak_db"] = slate.get("true_peak_db")

    return info


# --------------------------------------------------------------------------
# Loudness analysis via ffmpeg's ebur128 filter (full decode — expensive,
# hence gated behind checkboxes rather than always run).
# --------------------------------------------------------------------------
def analyze_loudness(filepath, want_integrated, want_peak, timeout=600):
    """Integrated loudness (I) and true peak via ffmpeg's ebur128 filter —
    a libebur128 implementation of EBU R128 / ITU-R BS.1770-4. This is the
    same underlying measurement Spotify, YouTube, Apple, Netflix, and
    broadcast (EBU R128 / ATSC A/85) all use; what differs between them is
    the *target* LUFS they normalize to (and, for Netflix, dialogue-gating
    instead of the full-programme gating used here — see
    PLATFORM_TARGETS / platform_delta_rows below)."""
    ffmpeg_path, _ = resolve_binaries()
    if not ffmpeg_path:
        return {}
    cmd = [
        ffmpeg_path,
        "-nostats",
        "-hide_banner",
        "-i",
        filepath,
        "-filter:a",
        "ebur128=peak=true",
        "-f",
        "null",
        "-",
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
        text = result.stderr
    except subprocess.TimeoutExpired:
        return {"lufs_integrated": None, "true_peak_db": None, "lra": None}

    out = {}
    if want_integrated:
        m = re.search(r"Integrated loudness:\s*\n\s*I:\s*(-?\d+\.?\d*)\s*LUFS", text)
        out["lufs_integrated"] = float(m.group(1)) if m else None
        m = re.search(r"Loudness range:\s*\n\s*LRA:\s*(-?\d+\.?\d*)\s*LU", text)
        out["lra"] = float(m.group(1)) if m else None
    if want_peak:
        m = re.search(r"True peak:\s*\n\s*Peak:\s*(-?\d+\.?\d*)\s*dBFS", text)
        out["true_peak_db"] = float(m.group(1)) if m else None
    return out


# --------------------------------------------------------------------------
# Platform loudness targets — all BS.1770-based measurements, they just
# normalize to different reference LUFS values (and Netflix dialogue-gates
# rather than measuring the full programme, which our ebur128 pass doesn't
# replicate — flagged in the comparison row rather than silently assumed).
# Sources: Spotify's own loudness docs (-14 LUFS / -1 dBTP via ITU 1770);
# EBU R128 (-23 LUFS); ATSC A/85 (-24 LKFS, US broadcast); Netflix delivery
# spec (-27 LKFS dialogue-gated, -2 dBTP ceiling); Apple Music Sound Check
# (-16 LUFS); YouTube (~-14 LUFS).
# --------------------------------------------------------------------------
PLATFORM_TARGETS = {
    "Spotify / YouTube / Amazon / Tidal": {"lufs": -14.0, "peak_dbtp": -1.0, "gated": "full-programme"},
    "Apple Music (Sound Check)": {"lufs": -16.0, "peak_dbtp": -1.0, "gated": "full-programme"},
    "EBU R128 (Broadcast EU)": {"lufs": -23.0, "peak_dbtp": -1.0, "gated": "full-programme"},
    "ATSC A/85 (US Broadcast)": {"lufs": -24.0, "peak_dbtp": -2.0, "gated": "full-programme"},
    "Netflix (delivery spec)": {"lufs": -27.0, "peak_dbtp": -2.0, "gated": "dialogue"},
}


def platform_delta_rows(info, target_key):
    """Extra (label, value) rows comparing measured LUFS/peak against a
    platform target. Only meaningful if lufs/true_peak were measured."""
    rows = []
    target = PLATFORM_TARGETS.get(target_key)
    if not target:
        return rows

    lufs = info.get("lufs_integrated")
    if lufs is not None:
        delta = lufs - target["lufs"]
        note = (
            " (dialogue-gated target — this is a full-programme measurement, treat as approximate)"
            if target["gated"] == "dialogue"
            else ""
        )
        rows.append((f"Δ vs {target_key}", f"{delta:+.1f} LU  (target {target['lufs']:.0f} LUFS){note}"))

    peak = info.get("true_peak_db")
    if peak is not None:
        over = peak - target["peak_dbtp"]
        status = "OK" if over <= 0 else f"exceeds ceiling by {over:.1f} dB"
        rows.append(
            (f"True Peak vs {target_key}", f"{peak:.1f} dBTP  (limit {target['peak_dbtp']:.0f} dBTP — {status})")
        )

    return rows


# --------------------------------------------------------------------------
# TVC slate-beep detection — heuristic: a .mov encoded in an "Animation"/
# QuickTime-family codec at ~22s or ~37s is assumed to be a TVC (TV
# commercial) master with a sync slate beep in its head. We isolate the
# first `window` seconds and measure true peak there specifically, since
# that's where the beep — not the actual spot content — lives.
#
# codec_name is generic per-family (ProRes is always "prores" regardless
# of Proxy/LT/422/HQ/4444/XQ — the profile lives in the separate `profile`
# field, and the fourcc lives in `codec_tag_string`, e.g. apch/ap4h/ap4x —
# never in codec_name). TVC_CODEC_WHITELIST matches on codec_name; use
# TVC_PRORES_PROFILES below if you want to narrow it to specific ProRes
# variants (e.g. only 422 HQ / 4444 masters) instead of any ProRes profile.
# --------------------------------------------------------------------------
TVC_CODEC_WHITELIST = {"qtrle", "prores", "rawvideo"}
TVC_PRORES_PROFILES = None  # e.g. {"422 hq", "4444"} to narrow; None = any ProRes profile
TVC_DURATIONS = (22.0, 37.0)
TVC_DURATION_TOLERANCE = 1.5  # seconds
TVC_SLATE_WINDOW = 7  # seconds from head to analyze


def is_probable_tvc(info, filepath):
    if not filepath.lower().endswith(".mov"):
        return False
    codec_name = (info.get("codec") or "").lower()
    if codec_name not in TVC_CODEC_WHITELIST:
        return False
    if codec_name == "prores" and TVC_PRORES_PROFILES is not None:
        profile = (info.get("codec_profile") or "").replace("ProRes ", "").lower()
        if profile not in TVC_PRORES_PROFILES:
            return False
    duration = info.get("duration")
    if duration is None:
        return False
    d = float(duration)
    return any(abs(d - t) < TVC_DURATION_TOLERANCE for t in TVC_DURATIONS)


def analyze_slate_beep(filepath, window=TVC_SLATE_WINDOW, timeout=120):
    """True peak within just the first `window` seconds (the slate/beep
    region), not the whole file. `-t` is placed before `-i` so ffmpeg
    stops reading input early instead of decoding the full commercial."""
    ffmpeg_path, _ = resolve_binaries()
    if not ffmpeg_path:
        return {}
    cmd = [
        ffmpeg_path,
        "-nostats",
        "-hide_banner",
        "-t",
        str(window),
        "-i",
        filepath,
        "-filter:a",
        "ebur128=peak=true",
        "-f",
        "null",
        "-",
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
        text = result.stderr
    except subprocess.TimeoutExpired:
        return {"true_peak_db": None}

    m = re.search(r"True peak:\s*\n\s*Peak:\s*(-?\d+\.?\d*)\s*dBFS", text)
    return {"true_peak_db": float(m.group(1)) if m else None}


# --------------------------------------------------------------------------
# Row builder — single source of truth for required + optional fields,
# used by the GUI results view and by every export format.
# --------------------------------------------------------------------------
def info_rows(info, options):
    rows = [
        (
            "Codec",
            info.get("codec", "N/A")
            + (f" · {info['codec_profile']}" if info.get("codec_profile") else "")
            + (f" ({info['codec_long']})" if info.get("codec_long") else ""),
        ),
        ("Dimensions", f"{info['width']}x{info['height']}" if info.get("width") else "N/A"),
        ("FPS", f"{info['fps']:.2f}" if info.get("fps") else "N/A"),
        ("Bitrate", human_bitrate(info.get("bitrate"))),
        ("Duration", human_duration(info.get("duration"), fps=info.get("fps"), frame_count=info.get("frame_count"))),
        ("Size", human_size(info.get("size_bytes"))),
    ]

    if options.get("audio_info"):
        if info.get("has_audio") and info.get("audio_codec"):
            rows.append(
                (
                    "Audio Track",
                    f"{info['audio_codec']} · {info.get('audio_channels', '?')}ch · "
                    f"{info.get('audio_sample_rate', '?')} Hz",
                )
            )
        else:
            rows.append(("Audio Track", "N/A (no audio track)"))

    if options.get("color_info"):
        rows.append(("Color / HDR", info.get("color_summary", "N/A")))

    if options.get("container_info"):
        rows.append(("Container", info.get("container", "N/A")))

    if options.get("creation_date"):
        rows.append(("Created", info.get("creation_date", "N/A")))

    if options.get("aspect_ratio"):
        rows.append(("Aspect Ratio", info.get("aspect_ratio", "N/A")))

    if options.get("lufs"):
        v = info.get("lufs_integrated")
        rows.append(("Integrated Loudness", f"{v:.1f} LUFS" if v is not None else "N/A (no audio track)"))
        lra = info.get("lra")
        if lra is not None:
            rows.append(("Loudness Range (LRA)", f"{lra:.1f} LU"))

    if options.get("true_peak"):
        v = info.get("true_peak_db")
        rows.append(("True Peak", f"{v:.1f} dBTP" if v is not None else "N/A (no audio track)"))

    platform_target = options.get("platform_target")
    if platform_target and platform_target != "None":
        rows.extend(platform_delta_rows(info, platform_target))

    if options.get("tvc_slate"):
        if info.get("is_tvc_candidate"):
            v = info.get("slate_beep_peak_db")
            rows.append(("Slate Beep Peak (first 7s)", f"{v:.1f} dBTP" if v is not None else "N/A"))
        else:
            rows.append(("Slate Beep Peak", "N/A (not a 22s/37s .mov match)"))

    return rows


# --------------------------------------------------------------------------
# Report builders
# --------------------------------------------------------------------------
def build_txt_report(scan_results, options, root_label):
    lines = [f"Media Scan Report — {root_label}", f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ""]
    for entry in scan_results:
        lines.append(f"[{entry['dir']}]  ({len(entry['files'])} file(s))")
        for filename, info in entry["files"]:
            lines.append(f"  {filename}")
            if "error" in info:
                lines.append(f"    ! {info['error']}")
                continue
            for label, value in info_rows(info, options):
                lines.append(f"    {label}: {value}")
        lines.append("")
    return "\n".join(lines)


def build_markdown_report(scan_results, options, root_label):
    lines = [
        f"# 🎞️ Media Scan Report",
        "",
        f"**Source:** `{root_label}`  ",
        f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "",
    ]
    for entry in scan_results:
        count = len(entry["files"])
        lines.append(f"## 📁 {entry['dir']} ({count} file{'s' if count != 1 else ''})")
        lines.append("")
        for filename, info in entry["files"]:
            icon = "🎬" if info.get("kind") == "video" else "🖼️"
            lines.append(f"### {icon} `{filename}`")
            lines.append("")
            if "error" in info:
                lines.append(f"> ⚠ {info['error']}")
                lines.append("")
                continue
            lines.append("| Field | Value |")
            lines.append("|---|---|")
            for label, value in info_rows(info, options):
                lines.append(f"| {label} | {str(value).replace('|', chr(92) + '|')} |")
            lines.append("")
    return "\n".join(lines)


_COLOR_MAP_HTML = {
    "label": "#2e7d32",
    "error": "#c62828",
    "dim": "#6b7280",
}

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{title}</title>
<style>
  body {{ background:#1e1f29; margin:0; padding:2rem; font-family:'SF Mono','Fira Code',Consolas,monospace; }}
  .wrap {{ max-width:900px; margin:0 auto; }}
  .card {{ background:#282a36; color:#f8f8f2; border-radius:10px; padding:1.25rem 1.5rem; margin-bottom:1rem; box-shadow:0 6px 20px rgba(0,0,0,.35); }}
  h1 {{ color:#8be9fd; }} h2 {{ color:#f1fa8c; }} h3 {{ color:#ff79c6; margin-bottom:.4rem; }}
  table {{ border-collapse:collapse; width:100%; margin-top:.4rem; }}
  td {{ padding:4px 10px; border-bottom:1px solid #44475a; }}
  td:first-child {{ color:#50fa7b; font-weight:bold; width:35%; }}
  .meta {{ color:#6272a4; font-size:12px; text-align:right; }}
  .err {{ color:#ff5555; }}
</style>
</head>
<body><div class="wrap">
<h1>🎞 Media Scan Report</h1>
<div class="meta">{meta}</div>
{body}
</div></body></html>
"""


def build_html_report(scan_results, options, root_label):
    body_parts = []
    for entry in scan_results:
        count = len(entry["files"])
        body_parts.append(
            f'<div class="card"><h2>📁 {html.escape(entry["dir"])} ' f'({count} file{"s" if count != 1 else ""})</h2>'
        )
        for filename, info in entry["files"]:
            icon = "🎬" if info.get("kind") == "video" else "🖼️"
            body_parts.append(f"<h3>{icon} {html.escape(filename)}</h3>")
            if "error" in info:
                body_parts.append(f'<p class="err">⚠ {html.escape(info["error"])}</p>')
                continue
            rows_html = "".join(
                f"<tr><td>{html.escape(label)}</td><td>{html.escape(str(value))}</td></tr>"
                for label, value in info_rows(info, options)
            )
            body_parts.append(f"<table>{rows_html}</table>")
        body_parts.append("</div>")

    return HTML_TEMPLATE.format(
        title=f"Media Scan Report — {html.escape(root_label)}",
        meta=f"Source: {html.escape(root_label)} — Generated {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        body="".join(body_parts),
    )


def build_docx_report(scan_results, options, root_label, output_path):
    from docx import Document
    from docx.shared import RGBColor

    GREEN = RGBColor(0x1E, 0x7A, 0x34)
    RED = RGBColor(0xC0, 0x1C, 0x28)

    doc = Document()
    doc.add_heading("🎞 Media Scan Report", level=0)
    p = doc.add_paragraph()
    p.add_run("Source: ").bold = True
    p.add_run(root_label)
    p = doc.add_paragraph()
    p.add_run("Generated: ").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    for entry in scan_results:
        count = len(entry["files"])
        doc.add_heading(f"📁 {entry['dir']}  ({count} file{'s' if count != 1 else ''})", level=1)
        for filename, info in entry["files"]:
            icon = "🎬" if info.get("kind") == "video" else "🖼️"
            doc.add_heading(f"{icon} {filename}", level=2)
            if "error" in info:
                run = doc.add_paragraph().add_run(f"⚠ {info['error']}")
                run.font.color.rgb = RED
                continue
            table = doc.add_table(rows=0, cols=2)
            table.style = "Light List Accent 1"
            for label, value in info_rows(info, options):
                cells = table.add_row().cells
                run = cells[0].paragraphs[0].add_run(label)
                run.bold = True
                run.font.color.rgb = GREEN
                cells[1].text = str(value)
            doc.add_paragraph()

    doc.save(output_path)


def save_report(fmt, output_path, scan_results, options, root_label):
    if fmt == "html":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(build_html_report(scan_results, options, root_label))
    elif fmt == "md":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(build_markdown_report(scan_results, options, root_label))
    elif fmt == "txt":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(build_txt_report(scan_results, options, root_label))
    elif fmt == "docx":
        build_docx_report(scan_results, options, root_label, output_path)
    else:
        raise ValueError(f"Unknown format: {fmt}")
